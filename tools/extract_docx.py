import sys
from docx import Document

def extract(path):
    doc = Document(path)
    texts = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            texts.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    texts.append(cell.text)
    print('\n'.join(texts))

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: extract_docx.py <path>')
        sys.exit(1)
    extract(sys.argv[1])
